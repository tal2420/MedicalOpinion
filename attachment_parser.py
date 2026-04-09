"""
Extract text from attached files (PDF, Word, plain text).
The extracted text is fed into email_parser for structured data extraction.

Handles a common issue with Hebrew PDFs: text is often extracted in reversed
character order (visual RTL stored as LTR bytes). This module detects and
corrects reversed Hebrew text automatically.
"""

import os
import io
import re


# ============================================================
# Hebrew text reversal detection and correction
# ============================================================

# Common Hebrew words that appear in medical/legal documents
# If these appear reversed in the text, the text needs reversing
_NORMAL_HEBREW_MARKERS = [
    "שם", "תאריך", "כתובת", "טלפון", "לידה", "ביקור", "רופא",
    "תובע", "נפגע", "מטופל", "חולה", "תעודת", "זהות", "משפחה",
    "סיכום", "אורולוגיה", "בדיקה", "ניתוח", "אבחנה", "טיפול",
    "מרפאה", "בית", "חולים", "רמבם", "מחלקה",
]

_REVERSED_HEBREW_MARKERS = [w[::-1] for w in _NORMAL_HEBREW_MARKERS]


def _is_reversed_hebrew(text):
    """Detect if Hebrew text has reversed character order.

    Checks for reversed versions of common Hebrew words.
    """
    if not text or len(text) < 20:
        return False

    # Count occurrences of normal vs reversed markers
    normal_count = 0
    reversed_count = 0

    for marker in _NORMAL_HEBREW_MARKERS:
        if marker in text:
            normal_count += 1

    for marker in _REVERSED_HEBREW_MARKERS:
        if marker in text:
            reversed_count += 1

    # If significantly more reversed markers found, text is reversed
    return reversed_count > normal_count and reversed_count >= 2


def _reverse_hebrew_text(text):
    """Reverse Hebrew text extracted in wrong character order from PDFs.

    Strategy: split each line into segments (Hebrew vs non-Hebrew),
    reverse the segment order, and reverse only Hebrew segments' characters.
    Non-Hebrew segments (numbers, Latin, punctuation) keep their internal order.
    """
    lines = text.split('\n')
    fixed_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            fixed_lines.append("")
            continue

        # Split line into segments: Hebrew chars vs everything else
        # Hebrew Unicode range: \u0590-\u05FF
        segments = re.findall(r'[\u0590-\u05FF]+|[^\u0590-\u05FF]+', line)

        # Reverse segment order (RTL reordering)
        segments.reverse()

        # Reverse characters only within Hebrew segments
        fixed_segments = []
        for seg in segments:
            if re.match(r'^[\u0590-\u05FF]+$', seg):
                fixed_segments.append(seg[::-1])
            else:
                fixed_segments.append(seg)

        fixed_lines.append(''.join(fixed_segments))

    return '\n'.join(fixed_lines)


def extract_text_from_bytes(filename, content_bytes):
    """Extract text from file content based on file extension.

    Args:
        filename: Original filename (used to determine type)
        content_bytes: Raw bytes of the file

    Returns:
        Extracted text as string, or empty string on failure.
    """
    if not content_bytes or not filename:
        return ""

    ext = os.path.splitext(filename.lower())[1]

    try:
        if ext == ".pdf":
            return _extract_pdf(content_bytes)
        elif ext in (".docx", ".doc"):
            return _extract_docx(content_bytes)
        elif ext in (".txt", ".text", ".csv"):
            return _extract_text(content_bytes)
        elif ext in (".rtf",):
            return _extract_rtf(content_bytes)
        elif ext in (".html", ".htm"):
            return _extract_html(content_bytes)
    except Exception as e:
        print(f"[attachment_parser] Error extracting {filename}: {e}")

    return ""


def extract_text_from_file(filepath):
    """Extract text from a file on disk.

    Args:
        filepath: Path to the file

    Returns:
        Extracted text as string.
    """
    if not os.path.exists(filepath):
        return ""

    with open(filepath, "rb") as f:
        content = f.read()

    return extract_text_from_bytes(os.path.basename(filepath), content)


def extract_all_attachments_text(attachments):
    """Extract text from a list of attachment dicts.

    Args:
        attachments: List of dicts with 'filename' and 'content' (bytes) keys

    Returns:
        Combined text from all attachments, with separators.
    """
    texts = []
    for att in attachments:
        filename = att.get("filename", "")
        content = att.get("content")
        if not content:
            continue

        text = extract_text_from_bytes(filename, content)
        if text and text.strip():
            # Add filename header for context
            texts.append(f"\n--- [{filename}] ---\n{text}")

    return "\n".join(texts)


def extract_all_from_folder(folder_path):
    """Extract text from all files in a folder (attachments on disk).

    Args:
        folder_path: Path to the מצורפים folder

    Returns:
        Combined text from all files.
    """
    if not os.path.exists(folder_path):
        return ""

    texts = []
    supported_exts = {".pdf", ".docx", ".doc", ".txt", ".text", ".csv", ".rtf", ".html", ".htm"}

    for filename in os.listdir(folder_path):
        ext = os.path.splitext(filename.lower())[1]
        if ext not in supported_exts:
            continue

        filepath = os.path.join(folder_path, filename)
        if not os.path.isfile(filepath):
            continue

        text = extract_text_from_file(filepath)
        if text and text.strip():
            texts.append(f"\n--- [{filename}] ---\n{text}")

    return "\n".join(texts)


# ============================================================
# Format-specific extractors
# ============================================================

def _extract_pdf(content_bytes):
    """Extract text from PDF bytes using pdfplumber.

    Automatically detects and corrects reversed Hebrew text,
    which is a common issue with Hebrew PDFs.
    """
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            # Limit to first 20 pages to avoid huge documents
            if i >= 20:
                text_parts.append("[... עמודים נוספים לא נסרקו ...]")
                break

            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            # Also try extracting tables (common in medical/legal docs)
            try:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            cells = [str(cell).strip() for cell in row if cell]
                            if cells:
                                text_parts.append(" | ".join(cells))
            except Exception:
                pass

    raw_text = "\n".join(text_parts)

    # Auto-detect and fix reversed Hebrew text
    if raw_text and _is_reversed_hebrew(raw_text):
        raw_text = _reverse_hebrew_text(raw_text)

    return raw_text


def _extract_docx(content_bytes):
    """Extract text from Word .docx bytes."""
    from docx import Document

    doc = Document(io.BytesIO(content_bytes))
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables (very common in legal/medical documents)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text_parts.append(" | ".join(cells))

    return "\n".join(text_parts)


def _extract_text(content_bytes):
    """Extract text from plain text files."""
    # Try UTF-8 first, then Windows Hebrew (cp1255), then Latin
    for encoding in ["utf-8-sig", "utf-8", "cp1255", "windows-1255", "iso-8859-8", "latin-1"]:
        try:
            return content_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content_bytes.decode("utf-8", errors="replace")


def _extract_rtf(content_bytes):
    """Basic RTF text extraction (strip RTF control codes)."""
    text = content_bytes.decode("utf-8", errors="replace")
    # Remove RTF control words
    text = re.sub(r'\\[a-z]+\d*\s?', ' ', text)
    # Remove braces
    text = re.sub(r'[{}]', '', text)
    # Clean up whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def _extract_html(content_bytes):
    """Extract text from HTML, stripping tags."""
    text = _extract_text(content_bytes)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&[a-z]+;', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()
