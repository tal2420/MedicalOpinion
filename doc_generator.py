import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import load_config, get_data_dir


def _create_template():
    """Create the default opinion Word template."""
    config = load_config()
    doc = Document()

    # Set default RTL and Hebrew font
    style = doc.styles["Normal"]
    style.font.name = "David"
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- Header ---
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(config.get("professor_name", "[שם הפרופסור]"))
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f'{config.get("professor_title", "פרופסור")} ל{config.get("professor_specialty", "אורולוגיה")}'
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if config.get("professor_license"):
        license_para = doc.add_paragraph()
        license_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = license_para.add_run(f'רישיון מס׳: {config["professor_license"]}')
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Contact info
    contact_parts = []
    if config.get("professor_phone"):
        contact_parts.append(f'טל: {config["professor_phone"]}')
    if config.get("professor_address"):
        contact_parts.append(config["professor_address"])
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact.add_run(" | ".join(contact_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Separator line
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_before = Pt(12)
    run = title.add_run("חוות דעת רפואית")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("")  # Spacer

    # --- Case details section ---
    details = [
        ("מס׳ תיק:", "{{case_number}}"),
        ("תאריך:", "{{date}}"),
        ("שם הנבדק/ת:", "{{plaintiff_name}}"),
        ("ת.ז.:", "{{plaintiff_id}}"),
        ("גורם מפנה:", "{{sender_name}}"),
        ("צד שכנגד:", "{{opposing_party}}"),
    ]

    for label, placeholder in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{label} ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(placeholder)
        run.font.size = Pt(12)

    doc.add_paragraph("")

    # --- Sections ---
    sections = [
        "רקע רפואי",
        "ממצאי הבדיקה",
        "סיכום ומסקנות",
        "הערכת נכות",
    ]

    for section_title in sections:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.space_before = Pt(12)
        run = p.add_run(section_title)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        content.add_run(f"{{{{{section_title}}}}}")

    # --- Signature ---
    doc.add_paragraph("")
    doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    sig = doc.add_paragraph()
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig.space_before = Pt(24)
    run = sig.add_run("בכבוד רב,")
    run.font.size = Pt(12)

    sig_name = doc.add_paragraph()
    sig_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_name.add_run(config.get("professor_name", "[שם הפרופסור]"))
    run.bold = True
    run.font.size = Pt(12)

    sig_title = doc.add_paragraph()
    sig_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_title.add_run(
        f'{config.get("professor_title", "פרופסור")} ל{config.get("professor_specialty", "אורולוגיה")}'
    )
    run.font.size = Pt(11)

    return doc


def get_template_path():
    """Get or create the template path."""
    path = os.path.join(get_data_dir(), "תבנית_חוו״ד.docx")
    if not os.path.exists(path):
        doc = _create_template()
        doc.save(path)
    return path


def generate_opinion(case_data, case_path):
    """Generate a Word opinion document from case data.

    Args:
        case_data: dict with case fields
        case_path: path to the case folder

    Returns:
        Path to the generated document.
    """
    template_path = get_template_path()
    doc = Document(template_path)

    # Replacement map
    replacements = {
        "{{case_number}}": str(case_data.get("case_number", "")),
        "{{date}}": datetime.now().strftime("%d/%m/%Y"),
        "{{plaintiff_name}}": str(case_data.get("plaintiff_name", "")),
        "{{plaintiff_id}}": str(case_data.get("plaintiff_id", "")),
        "{{sender_name}}": str(case_data.get("sender_name", "")),
        "{{opposing_party}}": str(case_data.get("opposing_party", "")),
    }

    # Section placeholders - leave empty for the professor to fill
    section_placeholders = ["רקע רפואי", "ממצאי הבדיקה", "סיכום ומסקנות", "הערכת נכות"]
    for section in section_placeholders:
        replacements[f"{{{{{section}}}}}"] = ""

    # Apply replacements
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Save
    opinions_dir = os.path.join(case_path, "חוות_דעת")
    os.makedirs(opinions_dir, exist_ok=True)

    plaintiff_name = case_data.get("plaintiff_name", "unnamed")
    filename = f"חוו״ד_{plaintiff_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    # Sanitize filename
    import re
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)

    output_path = os.path.join(opinions_dir, filename)
    doc.save(output_path)

    return output_path
